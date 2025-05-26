import os
import requests
from lxml import etree
from docx import Document


def get_resp_data():
    headers = {
        "accept": "text/html,application/xhtml+xml,application/xml;q=0.9,image/avif,image/webp,image/apng,*/*;q=0.8,application/signed-exchange;v=b3;q=0.7",
        "accept-language": "zh-CN,zh;q=0.9",
        "cache-control": "no-cache",
        "pragma": "no-cache",
        "priority": "u=0, i",
        "sec-ch-ua": "\"Chromium\";v=\"136\", \"Google Chrome\";v=\"136\", \"Not.A/Brand\";v=\"99\"",
        "sec-ch-ua-mobile": "?0",
        "sec-ch-ua-platform": "\"Windows\"",
        "sec-fetch-dest": "document",
        "sec-fetch-mode": "navigate",
        "sec-fetch-site": "none",
        "sec-fetch-user": "?1",
        "upgrade-insecure-requests": "1",
        "user-agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/136.0.0.0 Safari/537.36"
    }
    cookies = {
        "RK": "wBXAgeL+7Z",
        "ptcz": "d32de47cbb14984f7061648cd978a513ffe9489c31940393aba9e367d3a26e4f",
        "pac_uid": "0_TW95fim5cj3JB",
        "omgid": "0_TW95fim5cj3JB",
        "_qimei_uuid42": "1951900122f100c898afd5caee96148a5856032d21",
        "_qimei_fingerprint": "97252720a76ba3430321a6adbfa654e4",
        "_qimei_q36": "",
        "_qimei_h38": "8f9578ee98afd5caee96148a0200000bd19519",
        "ua_id": "omWiKOI02IRTVdD6AAAAAJYmTok1I5GvXht8mr0PszM=",
        "_clck": "1hirrgb|1|fw7|0",
        "wxuin": "48140577492014",
        "data_bizuin": "3950457871",
        "data_ticket": "rFipZCB8ueys1+V9FwMOrcUfZWaSvZAdI3ZJqSj2WF4MDyeAJVHEvzpfsCKHMoMD",
        "xid": "e585651ecd5aa8db833a906df20743c0",
        "mm_lang": "zh_CN",
        "rewardsn": "",
        "wxtokenkey": "777"
    }

    url = "https://mp.weixin.qq.com/s/L5KVENGPMcYsgQK9Jc6SUA"
    response = requests.get(url, headers=headers, cookies=cookies)
    response.encoding = 'utf-8'
    data = response.text

    html = etree.HTML(data)
    title = html.xpath('//*[@id="activity-name"]/text()')[0].strip()
    content_section = html.xpath('//div[@id="js_content"]/section') or html.xpath('//div[@id="js_content"]')
    content_section = content_section[0]
    elements = content_section.xpath('.//p | .//img[@class="rich_pages wxw-img"]')

    return title, elements


def save_image(img_url, index):
    folder = 'images'
    os.makedirs(folder, exist_ok=True)
    img_data = requests.get(img_url).content

    filename = f"图{index}.jpg"
    print(filename, '已保存')
    path = os.path.join(folder, filename)
    with open(path, 'wb') as f:
        f.write(img_data)
    return filename  # 仅返回名字用于文档标记


def build_doc(title, elements):
    doc = Document()
    doc.add_heading(title, level=1)
    image_index = 1

    for elem in elements:
        if elem.tag == 'p':
            text = ','.join(elem.xpath('.//text()')) if elem.xpath('.//text()') else '无文本'
            if text != '无文本':
                doc.add_paragraph(text)
        elif elem.tag == 'img':
            src = elem.get('data-src') or elem.get('src') or elem.get('data-original')
            if src:
                image_name = save_image(src, image_index)
                doc.add_paragraph(f"{image_name[:-4]}")  # 插入“图x”占位符，不插入图片本体
                image_index += 1

    output_path = 'output.docx'
    doc.save(output_path)
    print(f"✅ Word文档已生成：{output_path}")
    print(f"✅ 图片已保存到 ./images 目录中，共 {image_index - 1} 张")


def main():
    title, elements = get_resp_data()
    build_doc(title, elements)


if __name__ == '__main__':
    main()
