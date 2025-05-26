import os
import time

import requests
from lxml import etree
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn


def get_resp_data(url):
    headers = {
        "accept": "text/html,application/xhtml+xml,application/xml;q=0.9,image/avif,image/webp,image/apng,*/*;q=0.8,application/signed-exchange;v=b3;q=0.7",
        "accept-language": "zh-CN,zh;q=0.9",
        "cache-control": "no-cache",
        "pragma": "no-cache",
        "priority": "u=0, i",
        "sec-ch-ua": "\"Chromium\";v=\"136\", \"Google Chrome\";v=\"136\", \"Not.A/Brand\";v=\"99\"",
        "sec-ch-ua-mobile": "?0",
        "sec-ch-ua-platform": "\"Windows\"",
        "sec-fetch-dest": "document",
        "sec-fetch-mode": "navigate",
        "sec-fetch-site": "none",
        "sec-fetch-user": "?1",
        "upgrade-insecure-requests": "1",
        "user-agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/136.0.0.0 Safari/537.36"
    }
    cookies = {
        "RK": "wBXAgeL+7Z",
        "ptcz": "d32de47cbb14984f7061648cd978a513ffe9489c31940393aba9e367d3a26e4f",
        "pac_uid": "0_TW95fim5cj3JB",
        "omgid": "0_TW95fim5cj3JB",
        "_qimei_uuid42": "1951900122f100c898afd5caee96148a5856032d21",
        "_qimei_fingerprint": "97252720a76ba3430321a6adbfa654e4",
        "_qimei_q36": "",
        "_qimei_h38": "8f9578ee98afd5caee96148a0200000bd19519",
        "ua_id": "omWiKOI02IRTVdD6AAAAAJYmTok1I5GvXht8mr0PszM=",
        "_clck": "1hirrgb|1|fw7|0",
        "wxuin": "48140577492014",
        "data_bizuin": "3950457871",
        "data_ticket": "rFipZCB8ueys1+V9FwMOrcUfZWaSvZAdI3ZJqSj2WF4MDyeAJVHEvzpfsCKHMoMD",
        "xid": "e585651ecd5aa8db833a906df20743c0",
        "mm_lang": "zh_CN",
        "rewardsn": "",
        "wxtokenkey": "777"
    }

    response = requests.get(url, headers=headers, cookies=cookies)
    response.encoding = 'utf-8'
    data = response.text

    html = etree.HTML(data)
    title = html.xpath('//*[@id="activity-name"]/text()')[0].strip()
    content_section = html.xpath('//div[@id="js_content"]/section') or html.xpath('//div[@id="js_content"]')
    content_section = content_section[0]
    elements = content_section.xpath('.//p | .//img[@class="rich_pages wxw-img"]')

    return title, elements


def add_indent_paragraph(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(12)
    p.paragraph_format.first_line_indent = Pt(28)


def add_centered_text(doc, text):
    p = doc.add_paragraph()
    p.alignment = 1  # 居中
    run = p.add_run(text)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(12)


def build_doc(title, elements):
    # 创建以标题为名的文件夹
    base_folder = os.path.join(os.getcwd(), title)
    os.makedirs(base_folder, exist_ok=True)

    image_folder = os.path.join(base_folder, 'images')
    os.makedirs(image_folder, exist_ok=True)

    def save_image(img_url, index):
        img_data = requests.get(img_url).content
        filename = f"图{index}.jpg"
        print(filename, '保存')
        path = os.path.join(image_folder, filename)
        with open(path, 'wb') as f:
            f.write(img_data)
        return filename

    doc = Document()
    doc.add_heading(title, level=1)
    image_index = 1

    for elem in elements:
        if elem.tag == 'p':
            text = ''.join(elem.xpath('.//text()')).strip()
            if text:
                add_indent_paragraph(doc, text)
        elif elem.tag == 'img':
            src = elem.get('data-src') or elem.get('src') or elem.get('data-original')
            if src:
                save_image(src, image_index)
                add_centered_text(doc, f"图{image_index}")
                image_index += 1

    output_path = os.path.join(base_folder, f'{title}.docx')
    doc.save(output_path)
    print(f"✅ Word文档已生成：{output_path}")
    print(f"✅ 图片已保存到 {image_folder} 目录中，共 {image_index - 1} 张")


def main():
    wechat_links = [
        "https://mp.weixin.qq.com/s/TEJzG3dfY5Xia4J8f0KmGw",
        "https://mp.weixin.qq.com/s/ADzryNBIXIHD7vx_tkPgTA",
        "https://mp.weixin.qq.com/s/y0I_UDSN9aFfzYqN6uj16Q",
        "https://mp.weixin.qq.com/s/CNH14NeBntZvY3A-a4smkA",
        "https://mp.weixin.qq.com/s/geliGIs32e5p_Th8ioqIwA",
        "https://mp.weixin.qq.com/s/U3uOSzVICqxj6_OsdH_6Rg",
        "https://mp.weixin.qq.com/s/SVzgMpFnWwieJ5vierJxsA",
        "https://mp.weixin.qq.com/s/KdUY6EmQbXvvWn9PX0ppBA",
        "https://mp.weixin.qq.com/s/JoxfHUdvhPEGA1H6XqmYDw",
        "https://mp.weixin.qq.com/s/fAVck8iVPoLBhWg4XcBBqQ",
        "https://mp.weixin.qq.com/s/T_6zHuioyG97RNHRgPLzIw",
        "https://mp.weixin.qq.com/s/_XhnY-pUCLr2gyjzbwKcWA",
        "https://mp.weixin.qq.com/s/-qENCScxJkO_Szcqfe8WxA",
        "https://mp.weixin.qq.com/s/cea4zCf1gDsy5o7lz4NWjA",
        "https://mp.weixin.qq.com/s/0jrNPdhhJ0Iu8RY5O76XWQ",
        "https://mp.weixin.qq.com/s/fijk9tdE4CMuLC7kgdVvRQ",
        "https://mp.weixin.qq.com/s/Yo5ZFNFhDyYyDpnr68Hr8Q",
        "https://mp.weixin.qq.com/s/SVzgMpFnWwieJ5vierJxsA",
        "https://mp.weixin.qq.com/s/nEI_xFEyo0BkEK-V6m0S0g",
        "https://mp.weixin.qq.com/s/L5KVENGPMcYsgQK9Jc6SUA"
    ]
    for i in wechat_links:
        title, elements = get_resp_data(i)
        build_doc(title, elements)
        time.sleep(2)


if __name__ == '__main__':
    main()
