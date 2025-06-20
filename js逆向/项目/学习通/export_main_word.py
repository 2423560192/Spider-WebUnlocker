import docx
import requests
from lxml import etree
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

# 初始化 Word 文档
name = '现代史题库'
doc = Document()
doc.styles['Normal'].font.name = u'宋体'
doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
doc.styles['Normal'].font.size = Pt(12)
doc.add_heading(name + '导出', level=1)

headers = {
    "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,image/avif,image/webp,image/apng,*/*;q=0.8,application/signed-exchange;v=b3;q=0.7",
    "Accept-Language": "zh-CN,zh;q=0.9",
    "Cache-Control": "no-cache",
    "Connection": "keep-alive",
    "Pragma": "no-cache",
    "Sec-Fetch-Dest": "document",
    "Sec-Fetch-Mode": "navigate",
    "Sec-Fetch-Site": "none",
    "Sec-Fetch-User": "?1",
    "Upgrade-Insecure-Requests": "1",
    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/137.0.0.0 Safari/537.36",
    "sec-ch-ua": "\"Google Chrome\";v=\"137\", \"Chromium\";v=\"137\", \"Not/A)Brand\";v=\"24\"",
    "sec-ch-ua-mobile": "?0",
    "sec-ch-ua-platform": "\"Windows\""
}
cookies = {
    "fid": "10948",
    "k8s": "1750044485.218.18710.522447",
    "route": "2fe558bdb0a1aea656e6ca70ad0cad20",
    "jrose": "D7108BD6C96B461BE95552B387B9D46A.mooc-2281852096-lfnh5",
    "fanyamoocs": "11401F839C536D9E",
    "_uid": "350880339",
    "_d": "1750061424312",
    "UID": "350880339",
    "vc3": "UOxWdYo4wyNq85gsFu8iUyuvunQkKXrTGojZdhqDoF2kt%2BFiqZ%2F0lJn%2BalGtBY4ys1WP7r7%2BlNf2JoF3G76eho8IhR6awMPN6hNcdvaSksmpuHYOjsXcoqxUYTPmEQDdSc%2FE3a4hxk2mL%2Fcb7cQVg6JRaD8hHRg2XAx7%2BF%2FxqUE%3D2246faf87a8344e740926bc20d0b81ea",
    "uf": "b2d2c93beefa90dc1d23f9590714aa1c2cc7d7d5da0160ec2229a8c94184be7c8f37668ad3a0e3f985c671bf308e2004ad3044892db91af8c49d67c0c30ca5043ad701c8b4cc548c0234d89f51c3dccfae0d68279d81eefefb98ce0e6210c3884a878d0a9a7b05da6103a97f8cd189bcd4a6744322bbb14840d31641a6466c4fb68e4521633bc299d3aa8ff6b51b5e36da9735baa04d8d5fce71fc6e59483dd3b6b080ac54381e6829d55877572ae3296a8d01806d34576be9fdc681bdf07734",
    "cx_p_token": "e69e0526d0e74d7dcbf8a95be431b4d2",
    "p_auth_token": "eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9.eyJ1aWQiOiIzNTA4ODAzMzkiLCJsb2dpblRpbWUiOjE3NTAwNjE0MjQzMTMsImV4cCI6MTc1MDY2NjIyNH0.8gHADhDTt_dq8eqxnzJaxaXvwnDnMiIdcD4cKJNeTas",
    "xxtenc": "543b24dd06366c61c47c7c02ca589564",
    "DSSTASH_LOG": "C_38-UN_10404-US_350880339-T_1750061424313",
    "source": "\"\"",
    "spaceFid": "10948",
    "spaceRoleId": "\"\"",
    "tl": "1"
}
params = {
    "courseId": "253237693",
    "classId": "122999526",
    "cpi": "416702459",
    "workId": "44072813",
    "answerId": "54702781",
    "standardEnc": "87a0db588c39cdf051074133cfb2a0ce",
    "enc": "cf1319545f9bdb39bf308e95f306e77c"
}

# 请求并解析
response = requests.get('https://mooc1.chaoxing.com/mooc-ans/mooc2/work/dowork',
                        params=params, cookies=cookies, headers=headers)
html = etree.HTML(response.text)

ques_list = html.xpath('//div[@class="padBom50 questionLi fontLabel singleQuesId"]')
ans_list = html.xpath('//div[contains(@class, "stem_answer")]')

# 处理每道题
for idx, (q, a) in enumerate(zip(ques_list, ans_list), start=1):
    title = " ".join([
        i.strip() for i in q.xpath('./h3/text() | ./h3/following-sibling::p[position()<=2]/text() | ./h3/span/text()')
    ])
    para = doc.add_paragraph()
    para.add_run(f"{idx}. {title}").bold = True

    if '简答题' in title or '论述题' in title:
        answer = a.xpath('./div[@class="eidtDiv"]/textarea/text()')
        ans_text = answer[0].strip() if answer else "暂无答案"
        doc.add_paragraph(f"【参考答案】：{ans_text}")

    elif '材料题' in title:
        answer = a.xpath('.//div[@class="textDIV"]/textarea/text()')
        ans_text = answer[0].strip() if answer else "暂无答案"
        doc.add_paragraph(f"【参考答案】：{ans_text}")

    else:
        options = a.xpath('./div[@class="clearfix answerBg workTextWrap"]')
        correct = []
        for opt in options:
            key = opt.xpath('./span/text()')[0]
            val = opt.xpath('./div/p/text()')[0]
            is_correct = opt.xpath('./span[contains(@class, "check_answer")]')
            line = f"{key}. {val}"
            p = doc.add_paragraph(f"  {line}")
            if is_correct:
                correct.append(key)
                p.runs[0].bold = True
                p.runs[0].font.color.rgb = docx.shared.RGBColor(0, 128, 0)
        doc.add_paragraph(f"【正确答案】：{'、'.join(correct)}")

    doc.add_paragraph("")  # 空行

# 保存 Word
doc.save(name + ".docx")
print(f"✅ Word 文件已保存为：{name}.docx")
